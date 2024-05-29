from docx import Document

def save_errors_to_docx(errors, docx_path='/home/t24120/carkey_v1.0/backend/error_documentation.docx'):
    try:
        doc = Document(docx_path)
    except Exception as e:
        doc = Document()

    for error in errors:
        doc.add_paragraph(error)

    doc.save(docx_path)




